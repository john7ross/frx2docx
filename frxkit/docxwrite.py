# -*- coding: utf-8 -*-
"""Сетка блоков -> .docx (python-docx)."""

from __future__ import annotations

import base64
import re
from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor, Twips

from .common import (BORDER_STYLE_MAP, DXA_TO_EMU, UNIT_TO_DXA, split_lines)
from .layout import page_blocks, part_blocks

ALIGN_MAP = {
    "Left": WD_ALIGN_PARAGRAPH.LEFT,
    "Center": WD_ALIGN_PARAGRAPH.CENTER,
    "Right": WD_ALIGN_PARAGRAPH.RIGHT,
    "Justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
VALIGN_MAP = {"Top": "top", "Center": "center", "Bottom": "bottom"}
TEXT_DIRECTION = {90: "btLr", 270: "tbRl"}

# системные переменные FastReport, которым в Word есть поле
PAGE_FIELDS = {
    "[Page]": " PAGE ", "[Page#]": " PAGE ",
    "[TotalPages]": " NUMPAGES ", "[TotalPages#]": " NUMPAGES ",
}
FIELD_RE = re.compile(r"(\[(?:Page|TotalPages)#?\])")


# --------------------------------------------------------------------------
# низкоуровневые свойства ячеек
# --------------------------------------------------------------------------
def set_cell_borders(cell, brd, widths, style, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(old)
    borders = OxmlElement("w:tcBorders")
    val = BORDER_STYLE_MAP.get(style, "single")
    for side, key in (("top", "t"), ("left", "l"), ("bottom", "b"),
                      ("right", "r")):
        e = OxmlElement("w:" + side)
        if brd.get(key):
            e.set(qn("w:val"), val)
            e.set(qn("w:sz"), str(max(2, int(round(_width_of(widths, key) * 6)))))
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), color)
        else:
            e.set(qn("w:val"), "nil")
        borders.append(e)
    tc_pr.append(borders)


def _width_of(widths, key):
    if isinstance(widths, dict):
        return widths.get(key, 1.0) or 1.0
    return widths or 1.0


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_valign(cell, value):
    tc_pr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:vAlign")
    el.set(qn("w:val"), value)
    tc_pr.append(el)


def set_cell_direction(cell, angle):
    value = TEXT_DIRECTION.get(angle)
    if not value:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:textDirection")
    el.set(qn("w:val"), value)
    tc_pr.append(el)


def set_cell_margins(cell, left=40, right=40, top=0, bottom=0):
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:tcMar")):
        tc_pr.remove(old)
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom),
                      ("right", right)):
        e = OxmlElement("w:" + side)
        e.set(qn("w:w"), str(max(0, int(val))))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tc_pr.append(mar)


def set_row_height(row, twips):
    tr_pr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(max(1, int(twips))))
    h.set(qn("w:hRule"), "atLeast")
    tr_pr.append(h)


def set_table_fixed(table, widths_dxa):
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblLayout", "w:tblBorders"):
        for old in tbl_pr.findall(qn(tag)):
            tbl_pr.remove(old)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + side)
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tbl_pr.append(borders)
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "0")
    ind.set(qn("w:type"), "dxa")
    tbl_pr.append(ind)
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, w in zip(grid.findall(qn("w:gridCol")), widths_dxa):
            gc.set(qn("w:w"), str(int(w)))


# --------------------------------------------------------------------------
# наполнение ячейки
# --------------------------------------------------------------------------
def _apply_run(run, frag, base, obj):
    name = frag.get("face") or base["name"]
    run.font.name = name
    run.font.size = Pt(frag.get("size") or base["size"])
    run.font.bold = bool(frag.get("bold") or base["bold"])
    run.font.italic = bool(frag.get("italic") or base["italic"])
    run.font.underline = bool(frag.get("underline") or base["underline"])
    if frag.get("strike") or base["strike"]:
        run.font.strike = True
    if frag.get("sub"):
        run.font.subscript = True
    if frag.get("sup"):
        run.font.superscript = True
    color = frag.get("color") or obj.get("color")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def _add_field(paragraph, instr, frag, base, obj):
    """Поле Word (PAGE / NUMPAGES) вместо системной переменной FastReport."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    run = paragraph.add_run("1")
    _apply_run(run, frag, base, obj)
    fld.append(run._element)
    paragraph._p.append(fld)


def _add_hyperlink(paragraph, url, frag, base, obj):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = paragraph.add_run(frag.get("text", ""))
    styled = dict(frag)
    styled["underline"] = True
    styled["color"] = frag.get("color") or "0563C1"
    _apply_run(run, styled, base, obj)
    link.append(run._element)
    paragraph._p.append(link)


def _emit_fragment(paragraph, frag, base, obj):
    """Фрагмент текста: с полями Word вместо [Page]/[TotalPages]."""
    text = frag.get("text", "")
    if obj.get("url"):
        _add_hyperlink(paragraph, obj["url"], frag, base, obj)
        return
    if not FIELD_RE.search(text):
        run = paragraph.add_run(text)
        _apply_run(run, frag, base, obj)
        return
    for piece in FIELD_RE.split(text):
        if not piece:
            continue
        instr = PAGE_FIELDS.get(piece)
        if instr:
            _add_field(paragraph, instr, frag, base, obj)
        else:
            run = paragraph.add_run(piece)
            _apply_run(run, frag, base, obj)


def _cell_paddings(obj):
    """Padding объекта в dxa. По умолчанию — привычные 40/40."""
    pad = obj.get("pad")
    if not pad:
        return 40, 40, 0, 0
    left, top, right, bottom = pad
    return (int(round(left * UNIT_TO_DXA)), int(round(right * UNIT_TO_DXA)),
            int(round(top * UNIT_TO_DXA)), int(round(bottom * UNIT_TO_DXA)))


def fill_cell(cell, obj, warnings=None):
    cell.text = ""
    left, right, top, bottom = _cell_paddings(obj)
    if obj.get("image"):
        _insert_image(cell, obj, warnings)
        set_cell_valign(cell, VALIGN_MAP.get(obj["valign"], "top"))
        set_cell_margins(cell, left, right, top, bottom)
        set_cell_borders(cell, obj["brd"], obj["brd_w"],
                         obj["brd_style"], obj["brd_color"])
        if obj.get("fill") and obj["fill"] != "FFFFFF":
            set_cell_shading(cell, obj["fill"])
        return

    base = obj["font"]
    if obj.get("runs"):
        lines = obj["runs"]
    else:
        lines = [[{"text": ln, "bold": base["bold"], "italic": base["italic"],
                   "underline": base["underline"], "strike": base["strike"],
                   "size": None, "color": None, "face": None}]
                 for ln in split_lines(obj["text"])]

    first = cell.paragraphs[0]
    for i, frags in enumerate(lines):
        p = first if i == 0 else cell.add_paragraph()
        p.alignment = ALIGN_MAP.get(obj["align"], WD_ALIGN_PARAGRAPH.LEFT)
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        if obj.get("indent"):
            pf.first_line_indent = Twips(int(obj["indent"] * UNIT_TO_DXA))
        if obj.get("line_height"):
            pf.line_spacing = Twips(int(obj["line_height"] * UNIT_TO_DXA))
        set_mark_size(p, base["size"])
        if not frags:
            continue
        for frag in frags:
            _emit_fragment(p, frag, base, obj)

    set_cell_valign(cell, VALIGN_MAP.get(obj["valign"], "top"))
    set_cell_direction(cell, obj.get("angle", 0))
    set_cell_margins(cell, left, right, top, bottom)
    set_cell_borders(cell, obj["brd"], obj["brd_w"], obj["brd_style"],
                     obj["brd_color"])
    if obj["fill"] and obj["fill"] != "FFFFFF":
        set_cell_shading(cell, obj["fill"])


def _insert_image(cell, obj, warnings):
    """PictureObject хранит картинку как base64 в атрибуте Image."""
    try:
        blob = base64.b64decode(obj["image"], validate=False)
        para = cell.paragraphs[0]
        para.alignment = ALIGN_MAP.get(obj["align"], WD_ALIGN_PARAGRAPH.LEFT)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run()
        iw = obj.get("image_w") or obj["w"]
        ih = obj.get("image_h") or obj["h"]
        run.add_picture(BytesIO(blob),
                        width=Emu(int(iw * UNIT_TO_DXA * DXA_TO_EMU)),
                        height=Emu(int(ih * UNIT_TO_DXA * DXA_TO_EMU)))
    except Exception as exc:                              # noqa: BLE001
        if warnings is not None:
            warnings.append("не удалось вставить изображение %s: %s"
                            % (obj["name"], exc))


def blank_cell(cell):
    """Пустая ячейка не должна распирать строку: знак абзаца в ней —
    кегль 1 pt, иначе Word держит высоту строки по кеглю стиля Normal."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_mark_size(p, 1)
    set_cell_margins(cell, 0, 0)
    set_cell_borders(cell, {"t": 0, "r": 0, "b": 0, "l": 0}, 1.0, "Solid",
                     "000000")


def set_mark_size(paragraph, points):
    """Кегль знака абзаца (w:pPr/w:rPr/w:sz)."""
    p_pr = paragraph._p.get_or_add_pPr()
    r_pr = p_pr.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        p_pr.append(r_pr)
    for tag in ("w:sz", "w:szCs"):
        for old in r_pr.findall(qn(tag)):
            r_pr.remove(old)
        e = OxmlElement(tag)
        e.set(qn("w:val"), str(int(points * 2)))
        r_pr.append(e)


# --------------------------------------------------------------------------
# блоки
# --------------------------------------------------------------------------
def add_table(container, block, warnings=None):
    cols = block["cols"]
    widths = [int(round(w * UNIT_TO_DXA)) for w in cols]
    nrow, ncol = len(block["row_heights"]), len(cols)
    total = Emu(int(sum(widths) * DXA_TO_EMU))
    try:                          # Document.add_table(rows, cols, style)
        table = container.add_table(nrow, ncol)
    except TypeError:             # header/footer: add_table(rows, cols, width)
        table = container.add_table(nrow, ncol, total)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_fixed(table, widths)

    for r in range(nrow):
        set_row_height(table.rows[r], block["row_heights"][r] * UNIT_TO_DXA)
        for c in range(ncol):
            cell = table.cell(r, c)
            cell.width = Emu(int(widths[c] * DXA_TO_EMU))
            blank_cell(cell)

    covered = set()
    for ri, ci, rs, cs, obj in block["spans"]:
        if (ri, ci) in covered:
            if obj["text"].strip() and warnings is not None:
                warnings.append("ячейка %s попала в уже объединённую область "
                                "и пропущена" % (obj["name"] or obj["text"][:25]))
            continue
        for rr in range(ri, min(ri + rs, nrow)):
            for cc in range(ci, min(ci + cs, ncol)):
                covered.add((rr, cc))
        target = table.cell(ri, ci)
        if rs > 1 or cs > 1:
            r2 = min(ri + rs - 1, nrow - 1)
            c2 = min(ci + cs - 1, ncol - 1)
            try:
                target = table.cell(ri, ci).merge(table.cell(r2, c2))
            except Exception as exc:                      # noqa: BLE001
                if warnings is not None:
                    warnings.append("не удалось объединить ячейки (%s): %s"
                                    % (obj["name"], exc))
                target = table.cell(ri, ci)
        target.width = Emu(int(sum(widths[ci:ci + cs]) * DXA_TO_EMU))
        fill_cell(target, obj, warnings)
    return table


GAP_TWIPS = 20                 # высота технического абзаца между таблицами


def add_spacer(container, height_units):
    """Пустой промежуток шаблона. Сам абзац занимает GAP_TWIPS, поэтому
    на эту величину уменьшаем отступ — иначе пробелы накапливаются."""
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.line_spacing = Twips(GAP_TWIPS)
    pf.space_after = Twips(max(0, int(height_units * UNIT_TO_DXA) - GAP_TWIPS))
    set_mark_size(p, 1)
    p.add_run("").font.size = Pt(1)


def add_gap(container, twips=GAP_TWIPS):
    """Word склеивает две идущие подряд таблицы — разделяем их абзацем
    минимальной высоты."""
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Twips(twips)
    set_mark_size(p, 1)
    p.add_run("").font.size = Pt(1)


def write_blocks(container, blocks, warnings):
    for index, block in enumerate(blocks):
        if block["type"] == "spacer":
            add_spacer(container, block["h"])
            continue
        add_table(container, block, warnings)
        nxt = blocks[index + 1] if index + 1 < len(blocks) else None
        if nxt is None or nxt["type"] == "table":
            add_gap(container)


# --------------------------------------------------------------------------
# водяной знак
# --------------------------------------------------------------------------
# Водяной знак — повёрнутая VML-надпись в колонтитуле. Штатный для Word
# WordArt (v:textpath) в ряде версий не печатается, поэтому здесь обычный
# текст в v:rect без заливки и рамки: рисуется везде одинаково.
WATERMARK_XML = (
    '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    ' xmlns:v="urn:schemas-microsoft-com:vml"'
    ' xmlns:o="urn:schemas-microsoft-com:office:office"><w:r><w:pict>'
    '<v:rect id="frxWatermark%(n)d" o:spid="_x0000_s%(n)d" o:allowincell="f"'
    ' filled="f" stroked="f" style="position:absolute;margin-left:0;'
    'margin-top:0;width:%(w)dpt;height:%(h)dpt;rotation:315;'
    'z-index:-251657216;mso-position-horizontal:center;'
    'mso-position-horizontal-relative:margin;mso-position-vertical:center;'
    'mso-position-vertical-relative:margin">'
    '<v:textbox inset="0,0,0,0" style="mso-fit-shape-to-text:t">'
    '<w:txbxContent><w:p><w:pPr><w:jc w:val="center"/>'
    '<w:spacing w:before="0" w:after="0"/></w:pPr><w:r><w:rPr>'
    '<w:rFonts w:ascii="%(font)s" w:hAnsi="%(font)s" w:cs="%(font)s"/>'
    '<w:color w:val="C8C8C8"/><w:sz w:val="%(sz)d"/><w:szCs w:val="%(sz)d"/>'
    '</w:rPr><w:t xml:space="preserve">%(text)s</w:t></w:r></w:p>'
    '</w:txbxContent></v:textbox></v:rect></w:pict></w:r></w:p>')

_watermark_seq = [2049]


def add_watermark(section, text, font_name="Arial", size_pt=64):
    from docx.oxml import parse_xml
    escaped = (text.replace("&", "&amp;").replace("<", "&lt;")
                   .replace(">", "&gt;").replace('"', "&quot;"))
    _watermark_seq[0] += 1
    xml = WATERMARK_XML % {"w": max(220, int(len(text) * size_pt * 0.7)),
                           "h": int(size_pt * 1.6), "font": font_name,
                           "text": escaped, "sz": int(size_pt * 2),
                           "n": _watermark_seq[0]}
    header = section.header
    header.is_linked_to_previous = False
    header._element.insert(0, parse_xml(xml))


# --------------------------------------------------------------------------
# документ целиком
# --------------------------------------------------------------------------
def apply_geometry(section, page):
    width, height = page["width"], page["height"]
    margins = page["margins"]
    section.orientation = (WD_ORIENT.LANDSCAPE if width > height
                           else WD_ORIENT.PORTRAIT)
    section.page_width = Emu(width * DXA_TO_EMU)
    section.page_height = Emu(height * DXA_TO_EMU)
    section.left_margin = Emu(margins["left"] * DXA_TO_EMU)
    section.right_margin = Emu(margins["right"] * DXA_TO_EMU)
    section.top_margin = Emu(margins["top"] * DXA_TO_EMU)
    section.bottom_margin = Emu(margins["bottom"] * DXA_TO_EMU)
    if page.get("columns", 1) > 1:
        set_columns(section, page["columns"])


def set_columns(section, count):
    sect_pr = section._sectPr
    for old in sect_pr.findall(qn("w:cols")):
        sect_pr.remove(old)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), "425")
    sect_pr.append(cols)


def _fill_part(part, bands, page, warnings):
    """Колонтитул: таблицы вместо абзацев, первый пустой абзац убираем.
    Возвращает высоту колонтитула в units."""
    blocks, _ = part_blocks(bands, page["body_units"], warnings)
    if not blocks:
        return 0.0
    part.is_linked_to_previous = False
    write_blocks(part, blocks, warnings)
    for paragraph in part.paragraphs:
        if paragraph.text or paragraph._p.findall(".//" + qn("w:pict")):
            break                     # водяной знак и текст не трогаем
        paragraph._p.getparent().remove(paragraph._p)
        break
    return sum(b["h"] if b["type"] == "spacer" else sum(b["row_heights"])
               for b in blocks)


def _reserve_margins(section, page, head_units, foot_units):
    """FastReport ставит колонтитул вплотную к полю страницы и на его высоту
    ужимает тело. Word же кладёт колонтитул на расстоянии footer_distance от
    края, поэтому поле увеличиваем на высоту колонтитула, а расстояние делаем
    равным исходному полю — тогда тело получает ровно ту же высоту."""
    margins = page["margins"]
    if head_units > 0:
        section.header_distance = Emu(margins["top"] * DXA_TO_EMU)
        section.top_margin = Emu(int(margins["top"]
                                     + head_units * UNIT_TO_DXA) * DXA_TO_EMU)
    if foot_units > 0:
        section.footer_distance = Emu(margins["bottom"] * DXA_TO_EMU)
        section.bottom_margin = Emu(int(margins["bottom"]
                                        + foot_units * UNIT_TO_DXA) * DXA_TO_EMU)


def build_document(pages, warnings):
    """pages — список описаний из frxread.page_info()."""
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for index, page in enumerate(pages):
        section = (document.sections[0] if index == 0
                   else document.add_section(WD_SECTION.NEW_PAGE))
        apply_geometry(section, page)
        if page.get("watermark"):
            add_watermark(section, page["watermark"],
                          page.get("watermark_font", {}).get("name", "Arial"))
        head_h = _fill_part(section.header, page["header"], page, warnings)
        foot_h = _fill_part(section.footer, page["footer"], page, warnings)
        _reserve_margins(section, page, head_h, foot_h)
        blocks, _ = page_blocks(page, warnings)
        write_blocks(document, blocks, warnings)
    return document
