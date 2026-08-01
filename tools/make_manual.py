# -*- coding: utf-8 -*-
"""Собирает «КАК ПОЛЬЗОВАТЬСЯ.docx» — короткую инструкцию для тех, кто не
работает с командной строкой. Запуск:  python tools/make_manual.py
"""

from __future__ import annotations

import os
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
sys.path.insert(0, ROOT)

from docx import Document                                       # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH                   # noqa: E402
from docx.shared import Pt, RGBColor                            # noqa: E402

from frxkit import __version__                                  # noqa: E402

ACCENT = RGBColor(0x1F, 0x4E, 0x79)


def heading(doc, text, size=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    return p


def body(doc, text, bullet=False, mono=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bullet:
        p.paragraph_format.left_indent = Pt(18)
        text = "•  " + text
    run = p.add_run(text)
    run.font.size = Pt(11)
    if mono:
        run.font.name = "Consolas"
        run.font.size = Pt(10)
    return p


def build(path):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Конвертер шаблонов FastReport")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Инструкция для тех, кто не работает с командной "
                      "строкой.  Версия %s" % __version__)
    run.font.size = Pt(11)

    heading(doc, "Что это делает")
    body(doc, "Работает в обе стороны:")
    body(doc, "шаблон .frx  →  документ Word (.docx) и PDF", bullet=True)
    body(doc, "документ .docx, .pdf, .md или .txt  →  шаблон .frx", bullet=True)
    body(doc, "Нужно, чтобы шаблон можно было прочитать, показать юристам и "
              "согласовать правки, не открывая FastReport, — а потом вернуть "
              "поправленный текст обратно в шаблон.")
    body(doc, "Важно: это конвертер ШАБЛОНА, а не готового документа. "
              "Подстановки вида [root.FIO] остаются в тексте как есть — "
              "данные из базы не подставляются.")

    heading(doc, "Что нужно установить")
    body(doc, "Ничего. Python и все библиотеки уже лежат внутри папки. "
              "Требуется только Windows 64-бит.")
    body(doc, "PDF собирается через Microsoft Word, если он есть. Если Word "
              "и LibreOffice не найдутся, программа сделает PDF сама "
              "встроенным движком — результат будет почти таким же.")

    heading(doc, "Способ 1, самый простой")
    body(doc, "1. Скопируйте файлы в папку IN.")
    body(doc, "2. Дважды щёлкните КОНВЕРТИРОВАТЬ.bat и выберите формат "
              "для шаблонов: 1 — Word, 2 — PDF, 3 — оба.")
    body(doc, "3. Дождитесь надписи «Готово». Папка OUT откроется сама — "
              "в ней лежат готовые файлы.")
    body(doc, "Файлы .docx, .pdf, .md и .txt в папке IN превращаются в "
              "шаблоны .frx — для них вопрос про формат не задаётся.")

    heading(doc, "Способ 2, перетаскиванием")
    body(doc, "Выделите один или несколько файлов и перетащите их мышью прямо "
              "на КОНВЕРТИРОВАТЬ.bat. Программа спросит формат, готовые файлы "
              "появятся рядом с исходными.")

    heading(doc, "Проверить, что вёрстка не поехала")
    body(doc, "Дважды щёлкните «ПРОВЕРИТЬ ВЁРСТКУ.bat». Программа сверит, на "
              "каком расстоянии от края страницы стоит каждая надпись в "
              "шаблоне и в получившемся PDF.")
    body(doc, "Нормальный результат: в столбце «медиана» около 1 мм, в столбце "
              "«хвост» — ноль. Если в «хвосте» появились числа, программа "
              "покажет, какие именно надписи уехали и на сколько — эти места "
              "стоит посмотреть глазами.")

    heading(doc, "Если нужно отдать программу коллеге")
    body(doc, "Вариант А: скопируйте всю папку целиком. Она самодостаточна.")
    body(doc, "Вариант Б: дважды щёлкните «СОБРАТЬ EXE.bat» — получится один "
              "файл dist\\frx2docx.exe, который работает сам по себе. Для "
              "сборки нужен интернет, один раз.")

    heading(doc, "Что переносится")
    for line in ("текст, шрифты, начертание, размер, цвет, выравнивание;",
                 "таблицы, объединённые ячейки, рамки и заливка;",
                 "размер страницы, ориентация, поля, колонки, водяной знак;",
                 "верхние и нижние колонтитулы — с повтором на каждой "
                 "странице;",
                 "номера страниц [Page] и [TotalPages];",
                 "галочки, линии, прямоугольники, картинки;",
                 "штрихкоды и QR-коды, если в шаблоне записан постоянный "
                 "текст;",
                 "форматированный текст RichObject;",
                 "разметка вида <b>жирный</b> внутри ячеек."):
        body(doc, line, bullet=True)

    heading(doc, "Что НЕ переносится")
    for line in ("Скрипты на C# (ScriptText) не выполняются. Что скрипт "
                 "прячет, показывает или подставляет при печати — в результат "
                 "не попадёт.",
                 "Данные из базы. [root.Field] остаётся текстом.",
                 "Диаграммы, карты, подотчёты, градиентные заливки.",
                 "Страницы, помеченные Visible=false, пропускаются — как и при "
                 "печати из FastReport."):
        body(doc, line, bullet=True)
    body(doc, "О каждом таком случае программа пишет предупреждение с именем "
              "объекта. Молча ничего не теряется.")

    heading(doc, "Если что-то пошло не так")
    body(doc, "Запустите КОНВЕРТИРОВАТЬ.bat и прочитайте сообщения в чёрном "
              "окне: там написано, какой файл и почему не получился. "
              "Подробный разбор — в файлах README.ru.md и ARCHITECTURE.ru.md.")

    doc.save(path)
    return path


if __name__ == "__main__":
    out = os.path.join(ROOT, "КАК ПОЛЬЗОВАТЬСЯ.docx")
    print(build(out))
