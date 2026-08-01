# -*- coding: utf-8 -*-
"""Тесты frx2docx. Запуск:  python tests/test_frxkit.py

Внешних библиотек, кроме python-docx, не требуют. Тесты, которым нужны
системные шрифты или примеры из C:\\Scripts\\test, сами себя пропускают.
"""

from __future__ import annotations

import base64
import os
import shutil
import sys
import tempfile
import unittest

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
DATA = os.path.join(HERE, "data")
sys.path.insert(0, ROOT)

from frxkit import (barcode, common, doctree, fonts, frxread, frxwrite,  # noqa: E402
                    layout, pdfparse, pdfread, pdfwrite, rtf, textread)
from frxkit.cli import frx_to_docx, frx_to_pdf_direct, to_frx        # noqa: E402

FEATURES = os.path.join(DATA, "features.frx")
CONTRACT = os.path.join(DATA, "contract.md")
SUBSTITUTE = os.path.join(DATA, "substitute.frx")
# Папка с боевыми шаблонами: их в репозитории нет (в .frx лежит строка
# подключения к базе). Свою можно указать переменной окружения.
SAMPLES = os.environ.get("FRX2DOCX_SAMPLES", r"C:\Scripts\test")


class TempCase(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.mkdtemp(prefix="frxtest_")

    def tearDown(self):
        shutil.rmtree(self.tmp, ignore_errors=True)

    def out(self, name):
        return os.path.join(self.tmp, name)


# --------------------------------------------------------------------------
class TestCommon(unittest.TestCase):
    def test_colors(self):
        self.assertEqual(common.parse_color("Red"), "FF0000")
        self.assertEqual(common.parse_color("255, 0, 0"), "FF0000")
        self.assertEqual(common.parse_color("255, 255, 0, 0"), "FF0000")
        self.assertEqual(common.parse_color("#FF00FF00"), "00FF00")
        self.assertEqual(common.parse_color("#00FF00"), "00FF00")
        self.assertIsNone(common.parse_color("Transparent"))
        self.assertIsNone(common.parse_color(""))
        self.assertEqual(common.format_color("FF8000"), "255, 128, 0")

    def test_font(self):
        font = common.parse_font("Arial, 9pt, style=Bold, Italic")
        self.assertEqual(font["name"], "Arial")
        self.assertEqual(font["size"], 9.0)
        self.assertTrue(font["bold"] and font["italic"])
        self.assertFalse(font["underline"])
        again = common.parse_font(common.format_font(font))
        self.assertEqual(again, font)

    def test_border(self):
        self.assertEqual(common.parse_border("All"),
                         {"t": 1, "r": 1, "b": 1, "l": 1})
        self.assertEqual(common.parse_border("Left, Bottom"),
                         {"t": 0, "r": 0, "b": 1, "l": 1})
        self.assertEqual(common.format_border({"t": 1, "r": 1, "b": 1, "l": 1}),
                         "All")
        self.assertIsNone(common.format_border({"t": 0, "r": 0, "b": 0, "l": 0}))

    def test_padding(self):
        self.assertEqual(common.parse_padding("1, 2, 3, 4"), (1.0, 2.0, 3.0, 4.0))
        self.assertEqual(common.parse_padding(None), (2.0, 1.0, 2.0, 1.0))

    def test_inline_html(self):
        lines = common.parse_inline_html(
            "обычный <b>жирный</b><br><i>курсив</i>")
        self.assertEqual(len(lines), 2)
        self.assertEqual(lines[0][0]["text"], "обычный ")
        self.assertTrue(lines[0][1]["bold"])
        self.assertTrue(lines[1][0]["italic"])

    def test_html_roundtrip(self):
        source = "<b>жирный</b> и <i>курсив</i>"
        lines = common.parse_inline_html(source)
        again = common.parse_inline_html(common.runs_to_html(lines))
        self.assertEqual([f["text"] for f in again[0]],
                         [f["text"] for f in lines[0]])
        self.assertTrue(again[0][0]["bold"])


class TestRtf(unittest.TestCase):
    SAMPLE = ("{\\rtf1\\ansi\\ansicpg1251\\deff0"
              "{\\fonttbl{\\f0\\fnil\\fcharset204 Times New Roman;}}"
              "{\\colortbl ;\\red0\\green0\\blue192;}"
              "\\pard\\qc{\\f0\\fs20 обычный }{\\b\\f0\\fs20 жирный}"
              "{\\cf1\\f0\\fs28 крупный}\\par{\\f0\\fs20 вторая}}")

    def test_lines(self):
        lines, align = rtf.rtf_to_lines(self.SAMPLE)
        self.assertEqual(align, "Center")
        self.assertEqual(len(lines), 2)
        self.assertEqual(lines[0][0]["text"], "обычный ")
        self.assertTrue(lines[0][1]["bold"])
        self.assertEqual(lines[0][2]["size"], 14.0)
        self.assertEqual(lines[0][2]["color"], "0000C0")
        self.assertEqual(lines[1][0]["text"], "вторая")

    def test_text(self):
        self.assertEqual(rtf.rtf_to_text(self.SAMPLE),
                         "обычный жирныйкрупный\r\nвторая")

    def test_plain_passthrough(self):
        self.assertEqual(rtf.rtf_to_text("просто текст"), "просто текст")

    def test_roundtrip(self):
        lines, _ = rtf.rtf_to_lines(self.SAMPLE)
        again, _ = rtf.rtf_to_lines(rtf.lines_to_rtf(lines))
        self.assertEqual("".join(f["text"] for f in again[0]),
                         "обычный жирныйкрупный")
        self.assertTrue(again[0][1]["bold"])


class TestBarcode(unittest.TestCase):
    def test_ean13_checksum(self):
        pattern = barcode.ean13("460123456789")
        self.assertEqual(len(pattern), 95)
        self.assertTrue(pattern.startswith("101"))
        self.assertTrue(pattern.endswith("101"))

    def test_ean13_bad_length(self):
        with self.assertRaises(ValueError):
            barcode.ean13("123")

    def test_code128(self):
        pattern = barcode.code128("FRX-2026")
        self.assertTrue(set(pattern) <= {"0", "1"})
        self.assertTrue(pattern.endswith("1100011101011"))

    def test_render_png(self):
        png, square = barcode.render("Code128", "ABC123")
        self.assertIsNotNone(png)
        self.assertFalse(square)
        self.assertTrue(png.startswith(b"\x89PNG"))

    def test_render_qr(self):
        png, square = barcode.render("QR Code", "https://example.org")
        if png is None:
            self.skipTest("библиотека qrcode недоступна")
        self.assertTrue(square)
        self.assertTrue(png.startswith(b"\x89PNG"))

    def test_unknown(self):
        png, reason = barcode.render("Postnet", "12345")
        self.assertIsNone(png)
        self.assertIn("не поддерживается", reason)


class TestFonts(unittest.TestCase):
    def setUp(self):
        self.font = fonts.load("Arial")
        if self.font is None:
            self.skipTest("в системе нет ни одного подходящего шрифта")

    def test_metrics(self):
        self.assertGreater(self.font.units_per_em, 0)
        self.assertGreater(self.font.text_width("Договор", 10), 10)
        self.assertGreater(self.font.gid("Д"), 0)

    def test_wrap(self):
        lines = fonts.wrap("Заемщик обязуется возвратить сумму займа "
                           "и уплатить проценты", "Arial", 10, 120)
        self.assertGreater(len(lines), 1)
        self.assertTrue(all(len(x) for x in lines))

    def test_subset(self):
        gids = {self.font.gid(ch) for ch in "Договор 123"}
        blob = self.font.subset(gids)
        self.assertLess(len(blob), len(self.font.data))
        self.assertEqual(blob[:4], b"\x00\x01\x00\x00")

    def test_substitute(self):
        self.assertIsNotNone(fonts.load("Совершенно неизвестный шрифт"))

    def test_bundled_present(self):
        folder = fonts.bundled_dir()
        self.assertIsNotNone(folder, "папка fonts/ не найдена")
        names = [f.lower() for f in os.listdir(folder)]
        for want in ("liberationsans-regular.ttf", "liberationserif-bold.ttf",
                     "liberationmono-regular.ttf"):
            self.assertIn(want, names)
        self.assertTrue(any("license" in n for n in names),
                        "рядом со шрифтами нет файла лицензии")

    def test_bundled_fallback(self):
        """Неизвестное семейство подменяется запасным — с кириллицей."""
        font = fonts.load("PT Astra Serif")
        self.assertIsNotNone(font)
        self.assertIn("liberation", os.path.basename(font.path).lower())
        for ch in "Договор займа №":
            self.assertNotEqual(font.gid(ch), 0, "нет глифа для %r" % ch)

    def test_bundled_keeps_style(self):
        bold = fonts.load("PT Astra Serif", bold=True)
        italic = fonts.load("PT Astra Serif", italic=True)
        self.assertIn("bold", os.path.basename(bold.path).lower())
        self.assertIn("italic", os.path.basename(italic.path).lower())

    def test_fallback_metrics_match(self):
        """Liberation метрически совместимы: подмена не сдвигает переносы."""
        arial = fonts.load("Arial")
        spare = fonts.load("Совершенно неизвестный шрифт")
        if "liberation" not in os.path.basename(spare.path).lower():
            self.skipTest("запасной шрифт не Liberation")
        text = "Заемщик обязуется возвратить сумму займа"
        self.assertAlmostEqual(arial.text_width(text, 10),
                               spare.text_width(text, 10), places=2)

    def test_system_font_wins(self):
        """Системный шрифт имеет приоритет над запасным."""
        font = fonts.load("Arial")
        self.assertNotIn("liberation", os.path.basename(font.path).lower())


class TestFrxRead(unittest.TestCase):
    def setUp(self):
        self.root = frxread.read_frx(FEATURES)
        self.warnings = []
        self.pages = frxread.report_pages(self.root, self.warnings)

    def test_pages(self):
        self.assertEqual(len(self.pages), 2)

    def test_page_info(self):
        info = frxread.page_info(self.pages[0], self.warnings, set())
        self.assertEqual(info["watermark"], "ОБРАЗЕЦ")
        self.assertTrue(info["header"])
        self.assertTrue(info["footer"])
        names = {o["name"] for band in info["body"] for o in band["objs"]}
        self.assertIn("Title", names)
        self.assertIn("Line1", names)
        self.assertIn("Qr1", names)

    def test_barcode_rendered(self):
        info = frxread.page_info(self.pages[0], self.warnings, set())
        objs = {o["name"]: o for band in info["body"] for o in band["objs"]}
        self.assertIsNotNone(objs["Code128"]["image"])
        raw = base64.b64decode(objs["Code128"]["image"])
        self.assertTrue(raw.startswith(b"\x89PNG"))

    def test_rich_object(self):
        info = frxread.page_info(self.pages[0], self.warnings, set())
        objs = {o["name"]: o for band in info["body"] for o in band["objs"]}
        rich = objs["Rich1"]
        self.assertIsNotNone(rich["runs"])
        self.assertTrue(any(f["bold"] for line in rich["runs"] for f in line))

    def test_hidden_pages(self):
        root = frxread.read_frx(FEATURES)
        warn = []
        pages = frxread.report_pages(root, warn, include_hidden=True)
        self.assertEqual(len(pages), 2)

    def test_geometry(self):
        width, height, margins, body = frxread.page_geometry(self.pages[0], [])
        self.assertGreater(height, width)
        self.assertAlmostEqual(margins["left"] / common.MM_TO_DXA, 15.0, 1)
        self.assertGreater(body, 600)


class TestLayout(unittest.TestCase):
    def _rect(self, x, y, w, h, text=""):
        return {"name": text or "obj", "x": x, "y": y, "w": w, "h": h,
                "text": text, "brd": {"t": 0, "r": 0, "b": 0, "l": 0},
                "font": {"name": "Arial", "size": 9.0, "bold": False,
                         "italic": False, "underline": False, "strike": False},
                "color": None, "fill": None, "align": "Left", "valign": "Top",
                "brd_w": {"t": 1, "r": 1, "b": 1, "l": 1},
                "brd_style": "Solid", "brd_color": "000000",
                "pad": (2, 1, 2, 1), "image": None, "runs": None}

    def test_two_columns(self):
        band = {"name": "b", "top": 0, "objs": [
            self._rect(0, 0, 100, 20, "слева"),
            self._rect(100, 0, 100, 20, "справа")]}
        blocks = layout.gridify(band, 200, [])
        self.assertEqual(len(blocks), 1)
        self.assertEqual(len(blocks[0]["cols"]), 2)
        self.assertEqual(len(blocks[0]["spans"]), 2)

    def test_spacer(self):
        band = {"name": "b", "top": 0, "objs": [
            self._rect(0, 0, 200, 20, "верх"),
            self._rect(0, 100, 200, 20, "низ")]}
        blocks = layout.gridify(band, 200, [])
        self.assertEqual([b["type"] for b in blocks],
                         ["table", "spacer", "table"])

    def test_float_extraction(self):
        warn = []
        band = {"name": "b", "top": 0, "objs": [
            self._rect(0, 0, 200, 100, "фон"),
            self._rect(20, 20, 40, 20, "поверх")]}
        blocks = layout.gridify(band, 200, warn)
        self.assertTrue(any("поверх" in w for w in warn))
        self.assertEqual(len(blocks), 2)


class TestDocxWrite(TempCase):
    def test_build(self):
        from docx import Document
        path = self.out("features.docx")
        frx_to_docx(FEATURES, path, verbose=False)
        document = Document(path)
        text = "\n".join(p.text for p in document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
        self.assertIn("Проверка возможностей конвертера", text)
        self.assertIn("Объединение", text)
        self.assertEqual(len(document.sections), 2)
        header = document.sections[0].header
        self.assertFalse(header.is_linked_to_previous)

    def test_landscape_section(self):
        from docx import Document
        path = self.out("features.docx")
        frx_to_docx(FEATURES, path, verbose=False)
        document = Document(path)
        second = document.sections[1]
        self.assertGreater(second.page_width, second.page_height)


class TestPdfWrite(TempCase):
    def test_build_and_reparse(self):
        path = self.out("features.pdf")
        frx_to_pdf_direct(FEATURES, path, verbose=False)
        self.assertTrue(os.path.exists(path))
        doc = pdfparse.load(path)
        pages = doc.pages()
        self.assertEqual(len(pages), 2)
        warn = []
        reader = pdfread.ContentReader(doc, pages[0], warn)
        reader.run(doc.content(pages[0]))
        text = " ".join(i["text"] for i in reader.items)
        self.assertIn("Проверка", text)
        self.assertIn("Объединение", text)
        self.assertIn("стр. 1 из 2", text)

    def test_missing_fonts(self):
        """Шаблон на шрифтах, которых нет в Windows, в том числе с
        кириллическим именем семейства."""
        path = self.out("substitute.pdf")
        frx_to_pdf_direct(SUBSTITUTE, path, verbose=False)
        doc = pdfparse.load(path)
        warn = []
        page = doc.pages()[0]
        reader = pdfread.ContentReader(doc, page, warn).run(doc.content(page))
        text = " ".join(i["text"] for i in reader.items)
        self.assertIn("Заёмщик обязуется", text)
        self.assertIn("ЁЖИК", text)
        self.assertIn("1234-5678-9012", text)

    def test_missing_fonts_to_docx(self):
        path = self.out("substitute.docx")
        frx_to_docx(SUBSTITUTE, path, verbose=False)
        self.assertTrue(os.path.getsize(path) > 4000)

    def test_graphics(self):
        path = self.out("features.pdf")
        frx_to_pdf_direct(FEATURES, path, verbose=False)
        doc = pdfparse.load(path)
        warn = []
        reader = pdfread.ContentReader(doc, doc.pages()[0], warn)
        reader.run(doc.content(doc.pages()[0]))
        self.assertTrue(reader.lines)


class TestReverse(TempCase):
    def test_markdown(self):
        path = self.out("contract.frx")
        to_frx(CONTRACT, path, verbose=False)
        with open(path, encoding="utf-8") as fh:
            body = fh.read()
        self.assertIn("ДОГОВОР ЗАЙМА", body)
        self.assertIn("<TableObject", body)
        self.assertIn("<DataBand", body)
        root = frxread.read_frx(path)
        pages = frxread.report_pages(root, [])
        self.assertEqual(len(pages), 1)

    def test_markdown_then_docx(self):
        frx = self.out("contract.frx")
        docx_path = self.out("contract.docx")
        to_frx(CONTRACT, frx, verbose=False)
        frx_to_docx(frx, docx_path, verbose=False)
        from docx import Document
        text = "\n".join(p.text for p in Document(docx_path).paragraphs)
        for table in Document(docx_path).tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
        self.assertIn("Предмет договора", text)
        self.assertIn("[root.Summa]", text)

    def test_plain_text(self):
        src = self.out("note.txt")
        with open(src, "w", encoding="utf-8") as fh:
            fh.write("Первый абзац.\n\nВторой абзац.\n")
        path = self.out("note.frx")
        to_frx(src, path, verbose=False)
        with open(path, encoding="utf-8") as fh:
            body = fh.read()
        self.assertIn("Первый абзац.", body)
        self.assertIn("Второй абзац.", body)

    def test_docx_roundtrip(self):
        docx_path = self.out("features.docx")
        frx_to_docx(FEATURES, docx_path, verbose=False)
        frx_path = self.out("back.frx")
        to_frx(docx_path, frx_path, verbose=False)
        with open(frx_path, encoding="utf-8") as fh:
            body = fh.read()
        self.assertIn("Проверка возможностей конвертера", body)
        self.assertIn("Объединение", body)
        root = frxread.read_frx(frx_path)
        self.assertGreaterEqual(len(frxread.report_pages(root, [])), 1)

    def test_pdf_roundtrip(self):
        pdf_path = self.out("features.pdf")
        frx_to_pdf_direct(FEATURES, pdf_path, verbose=False)
        frx_path = self.out("back.frx")
        to_frx(pdf_path, frx_path, verbose=False)
        with open(frx_path, encoding="utf-8") as fh:
            body = fh.read()
        self.assertIn("Проверка возможностей конвертера", body)
        self.assertIn("Border.Lines", body)
        docx_path = self.out("back.docx")
        frx_to_docx(frx_path, docx_path, verbose=False)
        self.assertTrue(os.path.exists(docx_path))


class TestTextRead(unittest.TestCase):
    def test_inline(self):
        runs = textread.inline("обычный **жирный** и *курсив*", 11.0)
        self.assertEqual(len(runs), 4)
        self.assertTrue(runs[1]["bold"])
        self.assertTrue(runs[3]["italic"])

    def test_table(self):
        blocks = textread.parse_markdown(
            "| a | b |\n|---|---|\n| 1 | 2 |\n", [])
        tables = [b for b in blocks if b["type"] == "table"]
        self.assertEqual(len(tables), 1)
        self.assertEqual(len(tables[0]["rows"]), 2)

    def test_heading_and_list(self):
        blocks = textread.parse_markdown("# Заголовок\n\n- пункт\n", [])
        kinds = [b["type"] for b in blocks]
        self.assertEqual(kinds, ["p", "p"])
        self.assertTrue(blocks[0]["runs"][0]["bold"])
        self.assertTrue(blocks[1]["runs"][0]["text"].startswith("•"))


class TestPdfParse(TempCase):
    def test_filters(self):
        import zlib
        info = {"Filter": pdfparse.Name("FlateDecode")}
        self.assertEqual(
            pdfparse.decode_stream(info, zlib.compress(b"hello")), b"hello")
        info = {"Filter": pdfparse.Name("ASCIIHexDecode")}
        self.assertEqual(pdfparse.decode_stream(info, b"48656C6C6F>"), b"Hello")

    def test_lexer(self):
        parser = pdfparse.Parser(b"<< /Type /Page /N 3 /Kids [1 0 R] >>")
        value = parser.value()
        self.assertEqual(str(value["Type"]), "Page")
        self.assertEqual(value["N"], 3)
        self.assertIsInstance(value["Kids"][0], pdfparse.Ref)

    def test_cmap(self):
        data = (b"begincmap 2 beginbfchar <0003> <0020>\n<0024> <0410>\n"
                b"endbfchar endcmap")
        table = pdfread.parse_cmap(data)
        self.assertEqual(table[3], " ")
        self.assertEqual(table[0x24], "А")

    def test_cmap_range(self):
        data = b"beginbfrange <0041> <0043> <0410> endbfrange"
        table = pdfread.parse_cmap(data)
        self.assertEqual(table[0x41], "А")
        self.assertEqual(table[0x43], "В")


class TestSamples(TempCase):
    """Прогон по настоящим шаблонам, если они есть на этой машине."""

    def setUp(self):
        super().setUp()
        if not os.path.isdir(SAMPLES):
            self.skipTest("папка с примерами %s не найдена" % SAMPLES)
        self.files = [os.path.join(SAMPLES, f) for f in sorted(os.listdir(SAMPLES))
                      if f.lower().endswith(".frx")]
        if not self.files:
            self.skipTest("в %s нет .frx" % SAMPLES)

    def test_all_to_docx(self):
        for src in self.files:
            name = os.path.splitext(os.path.basename(src))[0]
            frx_to_docx(src, self.out(name + ".docx"), verbose=False)
            self.assertTrue(os.path.getsize(self.out(name + ".docx")) > 4000)

    def test_all_to_pdf(self):
        for src in self.files:
            name = os.path.splitext(os.path.basename(src))[0]
            frx_to_pdf_direct(src, self.out(name + ".pdf"), verbose=False)
            with open(self.out(name + ".pdf"), "rb") as fh:
                self.assertEqual(fh.read(5), b"%PDF-")

    def test_roundtrip_first(self):
        src = self.files[0]
        docx_path = self.out("a.docx")
        frx_to_docx(src, docx_path, verbose=False)
        to_frx(docx_path, self.out("b.frx"), verbose=False)
        frx_to_docx(self.out("b.frx"), self.out("c.docx"), verbose=False)
        self.assertTrue(os.path.exists(self.out("c.docx")))


if __name__ == "__main__":
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")
        except (AttributeError, ValueError):
            pass
    unittest.main(verbosity=2)
